import gradio as gr
from ocr_processing import DocumentOCR
import os
import secrets
from dotenv import load_dotenv
import base64
import logging
import json
import openai
from mistralai.client import MistralClient
import zipfile
import shutil
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables for authentication only
load_dotenv()

USERNAME = os.getenv('USERNAME')
PASSWORD = os.getenv('PASSWORD')

# Global variables to store user-provided API keys - use an in-memory dict with no defaults
api_keys = {
    "Mistral": "",
    "OpenAI": ""
}

# Initialize OCR processor
ocr_processor = None
available_engines = ["Tesseract"]  # Tesseract is always available

# Load custom theme
from custom_theme import delite_theme

def initialize_ocr_processor():
    """Initialize or reinitialize the OCR processor with current API keys."""
    global ocr_processor, available_engines

    try:
        # Clear existing environment variables specific to API keys
        # Note: DocumentOCR now primarily uses keys during its own __init__.
        # This block might be less critical unless other parts rely on these env vars.
        env_vars_to_clear = [f"{engine.upper()}_API_KEY" for engine in api_keys.keys()]
        original_env = {var: os.environ.get(var) for var in env_vars_to_clear}

        for var in env_vars_to_clear:
            if var in os.environ:
                 os.environ.pop(var)

        # Temporarily set environment variables for the keys provided by the user
        for engine, key in api_keys.items():
            if key:
                env_var = f"{engine.upper()}_API_KEY"
                os.environ[env_var] = key
                logger.info(f"Temporarily setting {env_var} for initialization")

        # Initialize processor (which now also attempts to init models and fetch OpenAI models)
        logger.info("Initializing OCR processor with potentially updated API keys")
        ocr_processor = DocumentOCR()
        available_engines = ["Tesseract"] # Start with Tesseract

        # Check which engines initialized successfully based on the processor's state
        if ocr_processor.easyocr is not None:
            available_engines.append("EasyOCR")
            logger.info("EasyOCR engine is available")
        else:
            logger.info("EasyOCR engine is not available (initialization failed)")

        if ocr_processor.mistral is not None:
            available_engines.append("Mistral")
            logger.info("Mistral OCR engine is available")
        else:
            logger.info("Mistral OCR engine is not available")

        if ocr_processor.openai is not None:
            available_engines.append("OpenAI")
            logger.info("OpenAI OCR engine is available")
            # Log the fetched models (already logged in DocumentOCR init)
            # logger.info(f"Fetched OpenAI Models: {ocr_processor.available_openai_models}")
        else:
            logger.info("OpenAI OCR engine is not available")

        # Restore original environment variables that were temporarily set
        # This might be important if other processes outside this app rely on them.
        for var, original_value in original_env.items():
            if original_value is None:
                os.environ.pop(var, None)
            else:
                os.environ[var] = original_value
            logger.info(f"Restored original value for {var}")

    except Exception as e: # Catch potential errors during DocumentOCR() init
        logger.error(f"Error during OCR initialization: {str(e)}", exc_info=True)
        # Reset processor and engines on failure
        ocr_processor = None
        available_engines = ["Tesseract"] # Fallback
        logger.warning("OCR Processor initialization failed. Only Tesseract might be available if base init succeeds.")
        # Restore environment variables even on failure
        try:
            if 'original_env' in locals():
                for var, original_value in original_env.items():
                    if original_value is None:
                        os.environ.pop(var, None)
                    else:
                        os.environ[var] = original_value
        except Exception as restore_e:
            logger.error(f"Failed to restore environment variables after init error: {restore_e}")

# --- Helper function to clear output directory ---
def clear_output_directory(output_dir_path_str):
    """Removes all files and subdirectories within the specified directory."""
    if not output_dir_path_str:
        logger.warning("Output directory path is not set. Skipping cleanup.")
        return

    output_dir_path = Path(output_dir_path_str)
    logger.info(f"Attempting to clear output directory: {output_dir_path}")

    if output_dir_path.exists() and output_dir_path.is_dir():
        for item in output_dir_path.iterdir():
            try:
                if item.is_file() or item.is_symlink():
                    item.unlink()
                    logger.debug(f"Deleted file: {item}")
                elif item.is_dir():
                    shutil.rmtree(item)
                    logger.debug(f"Deleted directory: {item}")
            except Exception as e:
                logger.error(f"Error deleting item {item}: {e}", exc_info=True)
        logger.info(f"Successfully cleared output directory: {output_dir_path}")
    elif not output_dir_path.exists():
        logger.info(f"Output directory {output_dir_path} does not exist. No cleanup needed.")
    else:
        logger.warning(f"Path {output_dir_path} exists but is not a directory. Skipping cleanup.")

# Initialize on startup
initialize_ocr_processor()

# Helper function to generate the markdown text for available engines
def get_available_engines_markdown():
    # Base the markdown on the *currently available* engines list
    lines = []
    if "Tesseract" in available_engines:
        lines.append("- Tesseract (✅ always available)") # Assuming Tesseract init is robust
    else:
        lines.append("- Tesseract (❌ failed to initialize)")

    easyocr_status = '✅ available' if 'EasyOCR' in available_engines else '❌ not available'
    lines.append(f"- EasyOCR ({easyocr_status})")

    mistral_status = '✅ available' if 'Mistral' in available_engines else '❌ not available'
    lines.append(f"- Mistral ({mistral_status})")

    openai_status = '✅ available' if 'OpenAI' in available_engines else '❌ not available'
    lines.append(f"- OpenAI ({openai_status})")
    return "\n".join(lines)

def save_api_key(api_key, engine):
    """Validate and save API key, reinitialize, and return status message."""
    if not api_key:
        return f"❓ No API key provided for {engine}. Please enter a valid key."

    logger.info(f"Attempting to save API key for {engine}")

    # --- Key Validation Step --- #
    try:
        if engine == "OpenAI":
            try:
                temp_client = openai.OpenAI(api_key=api_key)
                temp_client.models.list() # Lightweight check
                logger.info(f"OpenAI API key validation successful for key ending in ...{api_key[-4:]}")
            except openai.AuthenticationError:
                logger.warning(f"OpenAI API key validation failed (AuthenticationError) for key ending in ...{api_key[-4:]}")
                return f"❌ Invalid OpenAI API Key. Authentication failed."
            except Exception as e:
                logger.error(f"OpenAI API key validation failed for key ending in ...{api_key[-4:]}: {e}")
                return f"❌ OpenAI key validation failed: {str(e)}"

        # For Mistral, we'll just save the key without validation
        # The initialization process will handle validation errors

        # --- Save Key and Re-initialize --- #
        logger.info(f"Saving API key for {engine}")
        global api_keys
        api_keys[engine] = api_key

        # Reinitialize with the new key
        initialize_ocr_processor() # This will re-check all engines based on current api_keys

        # Check if the specific engine we tried to enable is now available *after* init
        success = engine in available_engines
        if success:
             status_message = f"✅ {engine} API key saved and engine initialized."
             if engine == "OpenAI" and ocr_processor and ocr_processor.available_openai_models:
                  # Don't list all models, just confirm they loaded or mention the first one
                  status_message += f" Model list loaded (e.g., {ocr_processor.available_openai_models[0]})."
             elif engine == "OpenAI":
                  status_message += " Check logs for model loading details."
             return status_message
        else:
            # This case occurs when initialization fails
            logger.error(f"{engine} key saved but engine failed to initialize. Check logs for details.")
            if engine == "Mistral":
                # Don't clear the key for Mistral - just report the initialization failure
                return f"⚠️ Mistral API key saved, but engine initialization failed. Check server logs for details."
            else:
                # For other engines, clear the key
                api_keys[engine] = ""
                initialize_ocr_processor() # Re-initialize again to ensure state is clean
                return f"❌ Error: {engine} key failed post-initialization. Check logs."

    except Exception as e:
        # Catch any unexpected errors during the whole process
        logger.error(f"Error saving API key for {engine}: {str(e)}", exc_info=True)
        return f"❌ Unexpected error processing API key for {engine}: {str(e)}"

def clear_api_key(engine):
    """Clear the API key for the specified engine."""
    global api_keys
    if engine in api_keys:
        api_keys[engine] = ""
        initialize_ocr_processor() # Re-initialize to update engine list
        logger.info(f"API key for {engine} has been cleared")
        return f"⚠️ {engine} API key has been cleared."
    else:
        return f"❓ Engine {engine} not found or doesn't use API keys."

# Modify process_document function signature - now takes a list of files and state
def process_document(files, ocr_engine, selected_openai_model, current_results_state):
    """Process a list of documents, update state, and UI components."""
    logger.info(f"Received {len(files) if files else 0} file(s) for processing with engine {ocr_engine}")

    # --- Clear Output Directory Before Processing ---
    if ocr_processor and ocr_processor.output_dir:
        clear_output_directory(ocr_processor.output_dir)
    else:
        logger.warning("OCR processor or output directory not available, cannot clear output directory.")

    # Initialize or reset state for this batch
    # Keep existing results if you want cumulative processing, or reset like this:
    # current_results_state = {"text": {}, "images": {}}
    # For now, let's assume we process only the newly uploaded batch
    processed_results = {"text": {}, "images": {}}
    errors_occurred = False
    error_messages = []

    # --- Initial UI State --- #
    initial_updates = {
        result_selector: gr.Dropdown(choices=[], value=None, visible=False),
        md_output: "Processing...",
        image_output: gr.update(value=None, visible=False),
        download_format: gr.Radio(visible=False),
        download_selected_btn: gr.Button(visible=False),
        download_all_btn: gr.Button(visible=False),
        download_options_md: gr.update(visible=False),
        download_trigger_md: gr.update(visible=False),
        single_download_trigger: gr.update(value=None, visible=False),
        zip_download_trigger: gr.update(value=None, visible=False)
    }

    if not files:
        initial_updates[md_output] = "Error: No files uploaded."
        # Return initial updates dictionary and the unchanged state
        return (*initial_updates.values(), current_results_state) # Unpack values in order

    # --- Pre-processing Checks --- #
    if ocr_processor is None:
        initialize_ocr_processor() # Attempt recovery
        if ocr_processor is None:
            initial_updates[md_output] = "Error: OCR processor failed to initialize. Check server logs."
            return (*initial_updates.values(), current_results_state)
        logger.warning("OCR processor was None, attempted re-initialization.")

    if ocr_engine not in available_engines:
        initial_updates[md_output] = f"Error: {ocr_engine} OCR is not available. Check API keys or server logs."
        return (*initial_updates.values(), current_results_state)

    if ocr_engine == "OpenAI" and not selected_openai_model:
        if not (ocr_processor and ocr_processor.available_openai_models):
             initial_updates[md_output] = "Error: OpenAI engine selected, but no models could be loaded."
        else:
            initial_updates[md_output] = "Error: OpenAI engine selected, but no specific model chosen."
        return (*initial_updates.values(), current_results_state)

    # --- Process Each File --- #
    for file_obj in files:
        original_filename = Path(file_obj.name).name # Use pathlib for robust name extraction
        logger.info(f"Processing file: {original_filename}")
        try:
            # Environment variable handling assumed done within DocumentOCR init/methods now
            _ , image_paths, result_text = ocr_processor.process_document(
                file_obj, # Pass the whole Gradio file object
                ocr_engine,
                openai_model=selected_openai_model if ocr_engine == "OpenAI" else None
            )

            # Check for errors during individual file processing
            if result_text is not None and result_text.startswith("Error:"):
                logger.error(f"Error processing {original_filename}: {result_text}")
                errors_occurred = True
                error_messages.append(f"- {original_filename}: {result_text}")
                # Store placeholder or error info
                processed_results["text"][original_filename] = result_text
                processed_results["images"][original_filename] = image_paths or [] # Store paths even on error
            elif result_text is None:
                logger.warning(f"No text extracted from {original_filename}.")
                errors_occurred = True
                error_msg = f"- {original_filename}: Could not extract text."
                error_messages.append(error_msg)
                processed_results["text"][original_filename] = "Error: Could not extract text."
                processed_results["images"][original_filename] = image_paths or []
            else:
                # Success for this file
                logger.info(f"Successfully processed {original_filename}. Text length: {len(result_text)}")
                processed_results["text"][original_filename] = result_text
                processed_results["images"][original_filename] = image_paths or []

        except Exception as e:
            error_msg = f"Error processing file {original_filename}: {str(e)}"
            logger.error(error_msg, exc_info=True)
            errors_occurred = True
            error_messages.append(f"- {original_filename}: Processing failed unexpectedly. Check logs.")
            processed_results["text"][original_filename] = f"Error: Processing failed unexpectedly."
            processed_results["images"][original_filename] = [] # No images if critical error

    # --- Update UI Based on Results --- #
    final_md_output = ""
    final_image_output_update = gr.update(value=None, visible=False)
    final_dropdown_update = gr.Dropdown(choices=[], value=None, visible=False)
    final_download_format_update = gr.Radio(visible=True, interactive=True)
    final_dl_selected_btn_update = gr.Button(visible=True)
    final_dl_all_btn_update = gr.Button(visible=True)
    final_dl_options_md_update = gr.update(visible=False)
    final_dl_trigger_md_update = gr.update(visible=False)

    processed_filenames = list(processed_results["text"].keys())

    if not processed_filenames: # Should not happen if input `files` was not empty, but safety check
        final_md_output = "Error: No files were processed successfully."
        if error_messages:
             final_md_output += "\n\nErrors:\n" + "\n".join(error_messages)
        # Hide download options if no success
        final_dl_options_md_update = gr.update(visible=False)
        final_dl_trigger_md_update = gr.update(visible=False)
    else:
        first_filename = processed_filenames[0] # Original filename key
        # Generate dropdown choices: label shows ".md", value is the original filename
        dropdown_choices = [(f"{Path(fn).stem}.md", fn) for fn in processed_filenames]

        final_md_output = processed_results["text"][first_filename]

        # Get image paths for the first file
        first_file_images = processed_results["images"].get(first_filename, [])
        valid_display_paths = [p for p in first_file_images if p is not None and os.path.exists(p)]
        final_image_output_update = gr.update(value=valid_display_paths, visible=bool(valid_display_paths))

        # Update dropdown with (label, value) pairs
        final_dropdown_update = gr.Dropdown(
            choices=dropdown_choices,
            value=first_filename, # Use the original filename as the internal value
            label="Select Processed File to View/Download",
            interactive=True,
            visible=True
        )
        # Show download options on success
        final_dl_options_md_update = gr.update(visible=True)
        final_dl_trigger_md_update = gr.update(visible=True)

        # Prepend overall error summary if any occurred
        if errors_occurred:
            error_summary = f"**Warning:** Processing completed with errors for some files:\\n" + "\n".join(error_messages) + "\n\n---\n\n"
            final_md_output = error_summary + final_md_output

    # Return updates for all relevant components and the new state
    # Ensure download triggers are hidden initially after processing
    final_single_dl_trigger_update = gr.update(visible=False)
    final_zip_dl_trigger_update = gr.update(visible=False)

    return (
        final_dropdown_update,
        final_md_output,
        final_image_output_update,
        final_download_format_update,
        final_dl_selected_btn_update,
        final_dl_all_btn_update,
        final_dl_options_md_update,
        final_dl_trigger_md_update,
        final_single_dl_trigger_update,
        final_zip_dl_trigger_update,
        processed_results # The new state dictionary
    )

# Function to clear OCR tab fields - Needs update
def clear_ocr_fields(current_results_state): # Takes state to clear it
    """Clears the input and output fields in the OCR tab, including state."""
    logger.info("Clearing OCR fields and results state.")
    default_engine = "Tesseract" if "Tesseract" in available_engines else (available_engines[0] if available_engines else None)
    cleared_state = {"text": {}, "images": {}} # Reset state
    return {
        file_input: None,
        ocr_engine: default_engine,
        image_output: gr.update(value=None, visible=False),
        md_output: "Extracted Text will appear here",
        result_selector: gr.Dropdown(choices=[], value=None, visible=False, interactive=False),
        download_format: gr.Radio(value="txt", visible=False, interactive=False),
        download_selected_btn: gr.Button(visible=False),
        download_all_btn: gr.Button(visible=False),
        # Add updates for the markdown components to hide them
        download_options_md: gr.update(visible=False),
        download_trigger_md: gr.update(visible=False),
        single_download_trigger: gr.update(value=None, visible=False), # Hide trigger and clear value on clear
        zip_download_trigger: gr.update(value=None, visible=False),    # Hide trigger and clear value on clear
        processed_results_state: cleared_state # Return cleared state
    }

# --- Functions for Result Display and Download --- #

def display_selected_result(selected_filename, current_results_state):
    """Updates the markdown and image preview based on dropdown selection."""
    if not selected_filename or not current_results_state or selected_filename not in current_results_state["text"]:
        logger.warning(f"display_selected_result: Filename '{selected_filename}' not found in state.")
        return {
            md_output: "Error: Could not load result for selected file.",
            image_output: gr.update(value=None, visible=False)
            # No need to update download triggers here, they remain hidden unless download is clicked
        }

    text_result = current_results_state["text"][selected_filename]
    image_paths = current_results_state["images"].get(selected_filename, [])
    valid_display_paths = [p for p in image_paths if p is not None and os.path.exists(p)]

    logger.info(f"Displaying result for: {selected_filename}")
    return {
        md_output: text_result,
        image_output: gr.update(value=valid_display_paths, visible=bool(valid_display_paths))
        # No need to update download triggers here
    }

def download_selected_file(selected_filename, format_type, current_results_state):
    """Generates a file for the selected document and returns its path for download."""
    logger.info(f"Request to download '{selected_filename}' as '{format_type}'")
    if not selected_filename or selected_filename not in current_results_state["text"]:
        logger.error(f"Download failed: Filename '{selected_filename}' not found in results.")
        gr.Warning(f"Cannot download: Result for '{selected_filename}' not found.")
        # Return update to hide the trigger
        return gr.update(visible=False)

    result_text = current_results_state["text"][selected_filename]
    if result_text.startswith("Error:"):
        logger.warning(f"Attempting to download a file with processing errors: {selected_filename}")
        gr.Warning(f"Cannot download: '{selected_filename}' had processing errors.")
        # Return update to hide the trigger
        return gr.update(visible=False)

    try:
        download_path = ocr_processor.download_ocr_result(result_text, format_type, original_filename=selected_filename)

        if download_path and os.path.exists(download_path):
            logger.info(f"Prepared file for download: {download_path}")
            # Return update to show the trigger and provide the file path
            return gr.update(value=download_path, visible=True)
        else:
            logger.error(f"Failed to create download file for {selected_filename}. Path: {download_path}")
            gr.Error(f"Failed to create download file for {selected_filename}.")
            # Return update to hide the trigger
            return gr.update(visible=False)
    except Exception as e:
        logger.error(f"Error during download preparation for {selected_filename}: {e}", exc_info=True)
        gr.Error(f"Error creating download for {selected_filename}: {e}")
        # Return update to hide the trigger
        return gr.update(visible=False)

def download_all_files(format_type, current_results_state):
    """Generates files for all results, zips them, and returns the zip path."""
    logger.info(f"Request to download all results as '{format_type}' in a ZIP archive.")
    filenames = list(current_results_state["text"].keys())
    if not filenames:
        logger.warning("Download all aborted: No results found.")
        gr.Warning("No processed files available to download.")
        # Return update to hide the trigger
        return gr.update(visible=False)

    # Create a temporary directory for individual files
    temp_dir = Path(ocr_processor.output_dir) / f"temp_zip_{secrets.token_hex(4)}"
    zip_path = Path(ocr_processor.output_dir) / f"ocr_results_{secrets.token_hex(8)}.zip"
    try:
        temp_dir.mkdir(parents=True, exist_ok=True)
        files_to_zip = []
        files_with_errors = []

        for filename in filenames:
            result_text = current_results_state["text"][filename]
            if result_text.startswith("Error:"):
                logger.warning(f"Skipping file with error in zip: {filename}")
                files_with_errors.append(filename)
                continue # Skip files that had processing errors

            # Generate the file in the temporary directory
            try:
                # Use modified download_ocr_result to create the file inside temp_dir
                # Need to ensure download_ocr_result respects the output path (modify it?)
                # OR: Create file directly here
                base_name = Path(filename).stem
                output_filename = f"{base_name}.{format_type}"
                output_path = temp_dir / output_filename

                # Reuse logic from download_ocr_result (adapt as necessary)
                if format_type == 'txt':
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(result_text)
                elif format_type == 'md':
                     with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(result_text) # Assuming text is already markdown
                elif format_type == 'doc':
                     # Requires python-docx - check if installed/add dependency
                    try:
                        from docx import Document
                        document = Document()
                        document.add_paragraph(result_text)
                        document.save(output_path)
                    except ImportError:
                        logger.error("python-docx not installed. Cannot create .doc file.")
                        gr.Error("Cannot create .doc file: python-docx package is missing.")
                        # Fallback or skip?
                        continue # Skip this file for .doc format
                    except Exception as docx_e:
                        logger.error(f"Failed to create .doc for {filename}: {docx_e}", exc_info=True)
                        gr.Warning(f"Failed to create .doc file for {filename}.")
                        continue # Skip this file
                else:
                    logger.warning(f"Unsupported format '{format_type}' for file {filename}")
                    continue # Skip unsupported formats

                if output_path.exists():
                    files_to_zip.append(output_path)
                else:
                    logger.warning(f"File not created for zip: {output_path}")

            except Exception as file_e:
                logger.error(f"Error generating file {filename} for zip: {file_e}", exc_info=True)
                files_with_errors.append(filename)

        if not files_to_zip:
            logger.warning("No valid files were generated to include in the zip.")
            if files_with_errors:
                 gr.Warning(f"Could not create zip: All {len(files_with_errors)} file(s) had errors or could not be generated.")
            else:
                 gr.Warning("Could not create zip: No files to include.")
            return gr.update(visible=False)

        # Create the zip file
        with zipfile.ZipFile(zip_path, 'w') as zf:
            for file_path in files_to_zip:
                zf.write(file_path, arcname=file_path.name) # Use only filename inside zip

        logger.info(f"Created zip archive: {zip_path} with {len(files_to_zip)} file(s).")
        if files_with_errors:
             gr.Info(f"Zip created, but {len(files_with_errors)} file(s) were skipped due to errors: {', '.join(files_with_errors)}")

        # Return update to show the trigger and provide the zip path
        return gr.update(value=str(zip_path), visible=True)

    except Exception as e:
        logger.error(f"Error creating zip file: {e}", exc_info=True)
        gr.Error(f"Failed to create zip archive: {e}")
        # Return update to hide the trigger
        return gr.update(visible=False)
    finally:
        # Clean up the temporary directory
        if temp_dir.exists():
            try:
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as cleanup_e:
                logger.error(f"Error cleaning up temp directory {temp_dir}: {cleanup_e}")


# --- Functions for Modal Confirmation --- #

def show_confirmation():
    """Returns updates dictionary to show the confirmation UI elements."""
    return {
        clear_confirm_msg: gr.Markdown(value="⚠️ Results exist. Are you sure you want to clear everything?", visible=True),
        confirm_clear_btn: gr.Button(visible=True),
        cancel_clear_btn: gr.Button(visible=True)
    }

def hide_confirmation():
    """Returns updates dictionary to hide the confirmation UI elements."""
    return {
        clear_confirm_msg: gr.Markdown(value="", visible=False),
        confirm_clear_btn: gr.Button(visible=False),
        cancel_clear_btn: gr.Button(visible=False)
    }

# Modify handle_clear_click to check state
def handle_clear_click(current_results_state):
    """Handles the main Clear button click based on results state."""
    # Check if the results state actually contains processed files
    if current_results_state and current_results_state.get("text"):
        # Results exist, show confirmation
        updates = show_confirmation()
        # Add gr.update() for main fields to indicate no change yet
        # Need to list all potentially affected UI components from the OCR tab
        updates[file_input] = gr.update()
        updates[ocr_engine] = gr.update()
        updates[image_output] = gr.update()
        updates[md_output] = gr.update()
        updates[result_selector] = gr.update()
        updates[download_format] = gr.update()
        updates[download_selected_btn] = gr.update()
        updates[download_all_btn] = gr.update()
        updates[single_download_trigger] = gr.update() # Pass trigger update through
        updates[zip_download_trigger] = gr.update()    # Pass trigger update through
        updates[processed_results_state] = gr.update() # Pass state through
        # Add updates for the new markdown components
        updates[download_options_md] = gr.update()
        updates[download_trigger_md] = gr.update()

        return updates
    else:
        # No results, clear directly and ensure confirmation is hidden
        updates = clear_ocr_fields(current_results_state) # Clears fields and returns updates dict including cleared state
        updates.update(hide_confirmation()) # Adds updates for confirmation UI
        # Remove the state update from clear_ocr_fields dict as it's handled separately
        cleared_state = updates.pop(processed_results_state)
        # Get the trigger updates
        single_dl_trigger_update = updates.pop(single_download_trigger)
        zip_dl_trigger_update = updates.pop(zip_download_trigger)
        # Get the markdown updates
        dl_options_md_update = updates.pop(download_options_md)
        dl_trigger_md_update = updates.pop(download_trigger_md)

        # Ensure confirmation UI is hidden explicitly in the return tuple
        hide_confirm_updates = hide_confirmation()
        clear_confirm_msg_update = hide_confirm_updates[clear_confirm_msg]
        confirm_clear_btn_update = hide_confirm_updates[confirm_clear_btn]
        cancel_clear_btn_update = hide_confirm_updates[cancel_clear_btn]

        # Unpack the dictionary values and add the confirmation UI, triggers, markdown, and cleared state at the end
        return (
            *updates.values(), # Unpack remaining UI updates from clear_ocr_fields
            dl_options_md_update, dl_trigger_md_update, # Add markdown updates
            clear_confirm_msg_update, confirm_clear_btn_update, cancel_clear_btn_update, # Add confirmation UI updates
            single_dl_trigger_update, zip_dl_trigger_update, # Add trigger updates
            cleared_state      # Add cleared state
        )

# Modify clear_and_hide_confirmation to clear state and return a dictionary
def clear_and_hide_confirmation(current_results_state): # Takes state to clear it
    """Clears the main fields, state, hides the confirmation UI, and returns an updates dictionary."""
    logger.info("Confirm Clear clicked: Clearing fields and hiding confirmation.")

    # --- Clear Output Directory on Confirm Clear ---
    if ocr_processor and ocr_processor.output_dir:
        clear_output_directory(ocr_processor.output_dir)
    else:
        logger.warning("OCR processor or output directory not available during clear confirmation, cannot clear output directory.")

    # Get updates dictionary for clearing fields and state
    updates = clear_ocr_fields(current_results_state)
    # Get updates dictionary for hiding confirmation UI
    hide_updates = hide_confirmation()
    # Merge the two dictionaries
    updates.update(hide_updates)
    # Return the combined dictionary
    return updates

# Create Gradio interface
with gr.Blocks(theme=delite_theme) as demo:
    gr.Markdown("# Document OCR")
    gr.Markdown("Upload a document (PDF or image) to extract text using OCR.")

    # --- Add State to store results ---
    # Stores a dictionary mapping original filename to extracted text: {filename: text}
    # Also stores image paths per file: {filename: [path1, path2,...]}
    processed_results_state = gr.State({"text": {}, "images": {}})

    with gr.Tabs():
        # OCR Tab
        with gr.Tab("OCR"):
            with gr.Row():
                with gr.Column():
                    gr.Markdown("Upload Document(s) to OCR")
                    # --- Allow multiple files ---
                    file_input = gr.File(label="Upload Document(s)", file_count="multiple")
                    # Initialize Radio button with currently available engines
                    ocr_engine = gr.Radio(
                        choices=available_engines,
                        value=available_engines[0] if available_engines else None,
                        label="OCR Engine",
                        interactive=True
                    )
                    # Main action buttons
                    with gr.Row():
                        process_btn = gr.Button("Process Documents", variant="primary", scale=1) # Renamed slightly
                        clear_btn = gr.Button("Clear", variant="secondary", scale=1)

                    # Confirmation UI (initially hidden) - unchanged for now
                    clear_confirm_msg = gr.Markdown(value="", visible=False)
                    confirm_clear_btn = gr.Button("Confirm Clear", variant="stop", visible=False)
                    cancel_clear_btn = gr.Button("Cancel", variant="secondary", visible=False)
                    with gr.Row(): # Contains the confirmation message
                        clear_confirm_msg
                    with gr.Row(): # Contains the confirmation buttons
                        confirm_clear_btn
                        cancel_clear_btn

                    image_output = gr.Gallery(label="Document Pages Preview", visible=False) # Renamed, initially hidden

                with gr.Column():
                    # --- New Result Selection and Download UI ---
                    gr.Markdown("Extracted Text & Download")
                    result_selector = gr.Dropdown(
                        label="Select Processed File to View/Download",
                        choices=[],
                        value=None,
                        interactive=True,
                        visible=False # Initially hidden
                    )
                    md_output = gr.Markdown(label="Extracted Text", container=True, show_copy_button=True, value="Extracted Text will appear here")

                    # Assign variable and set initial visibility
                    download_options_md = gr.Markdown("Download Options", visible=False)
                    with gr.Row():
                        download_format = gr.Radio(
                            choices=["txt", "md", "doc"],
                            value="txt",
                            label="Download Format",
                            interactive=True,
                            scale=1,
                            visible=False # Initially hidden
                        )
                    with gr.Row():
                        download_selected_btn = gr.Button("Download Selected", variant="secondary", scale=1, visible=False) # Initially hidden
                        download_all_btn = gr.Button("Download All (ZIP)", variant="secondary", scale=1, visible=False) # Initially hidden

                    # --- Change: Make File components visible but non-interactive ---
                    # Assign variable and set initial visibility
                    download_trigger_md = gr.Markdown("Download Trigger Area (ignore)", visible=False) # Add explanation for user
                    single_download_trigger = gr.File(
                        label="Selected File Download",
                        visible=False, # Initially hidden
                        interactive=False
                    )
                    zip_download_trigger = gr.File(
                        label="ZIP Archive Download",
                        visible=False, # Initially hidden
                        interactive=False
                    )

        # API Keys Tab
        with gr.Tab("API Keys"):
            gr.Markdown("### Configure OCR API Keys")
            gr.Markdown("Enter your API keys for Mistral and OpenAI to enable their OCR engines. Keys are stored in memory only for the current session.")
            gr.Markdown("⚠️ **Security Note**: API keys are stored in memory and are not persisted when the server restarts.")

            with gr.Accordion("Mistral OCR", open=False):
                mistral_status = gr.Markdown(label="Status", value="")
                with gr.Column():
                    mistral_key = gr.Textbox(
                        label="Mistral API Key",
                        placeholder="Enter your Mistral API key...",
                        type="password",
                        value=""
                    )
                    with gr.Row():
                        mistral_save_btn = gr.Button("Save Mistral API Key", variant="primary")
                        mistral_clear_btn = gr.Button("Clear Key", variant="stop")

            with gr.Accordion("OpenAI OCR", open=False):
                openai_status = gr.Markdown(label="Status", value="")
                with gr.Column():
                    openai_key = gr.Textbox(
                        label="OpenAI API Key",
                        placeholder="Enter your OpenAI API key...",
                        type="password",
                        value=""  # Don't display key values
                    )
                    # Add the model selection dropdown here
                    openai_model_select = gr.Dropdown(
                        label="Select OpenAI Model",
                        choices=[], # Will be populated dynamically
                        value=None,
                        visible=False, # Initially hidden
                        interactive=True
                    )
                    with gr.Row():
                        openai_save_btn = gr.Button("Save OpenAI API Key", variant="primary")
                        openai_clear_btn = gr.Button("Clear Key", variant="stop")

            gr.Markdown("### Available OCR Engines")
            gr.Markdown("The following OCR engines are currently available:")
            # Update Markdown text using the helper function
            available_engines_text = gr.Markdown(value=get_available_engines_markdown())

    # --- Event Handlers --- #

    # Process button click - Needs significant changes
    # Inputs: file_input (list), ocr_engine, selected_openai_model, processed_results_state
    # Outputs: result_selector, md_output, image_output, download_format, download_selected_btn, download_all_btn, processed_results_state
    process_btn.click(
        fn=process_document,
        inputs=[file_input, ocr_engine, openai_model_select, processed_results_state], # Add state
        # Update outputs for new UI components and match the return tuple order
        outputs=[
            result_selector,      # 1
            md_output,            # 2
            image_output,         # 3
            download_format,      # 4
            download_selected_btn,# 5
            download_all_btn,     # 6
            download_options_md,  # 7 - Added
            download_trigger_md,  # 8 - Added
            single_download_trigger, # 9
            zip_download_trigger,    # 10
            processed_results_state # 11 - Pass back the updated state
        ]
    )

    # Update display when dropdown selection changes
    result_selector.change(
        fn=display_selected_result,
        inputs=[result_selector, processed_results_state],
        outputs=[md_output, image_output]
    )

    # Trigger single file download button click
    download_selected_btn.click(
        fn=download_selected_file,
        inputs=[result_selector, download_format, processed_results_state],
        outputs=[single_download_trigger] # Output to the hidden file component
    )

    # Trigger zip file download button click
    download_all_btn.click(
        fn=download_all_files,
        inputs=[download_format, processed_results_state],
        outputs=[zip_download_trigger] # Output to the hidden file component
    )

    # Define the comprehensive UI update function
    def update_all_ui_elements():
        """Updates engine list, radio choices, statuses, and OpenAI model dropdown based on current state."""
        # Update available engines markdown
        markdown_text = get_available_engines_markdown()

        # Update OCR engine radio button choices and value
        # Prefer Tesseract if available, otherwise the first in the list
        default_engine_choice = "Tesseract" if "Tesseract" in available_engines else (available_engines[0] if available_engines else None)
        radio_update = gr.Radio(
            choices=available_engines,
            value=default_engine_choice,
            label="OCR Engine",
            interactive=True
        )

        # Determine Mistral Status
        mistral_status_text = ""
        if "Mistral" in available_engines:
            mistral_status_text = "✅ Mistral Engine Available."
        elif "Mistral" in api_keys and api_keys["Mistral"]:
             mistral_status_text = "❌ Mistral Engine Failed to Initialize. Check API Key/Logs."
        else:
             mistral_status_text = "ℹ️ Enter Mistral API Key to enable."


        # Determine OpenAI Status & Model Dropdown State
        openai_models = []
        openai_dropdown_visible = False
        openai_default_model = None
        openai_status_text = ""

        if "OpenAI" in available_engines:
             # Engine is available, check for models from the processor instance
             if ocr_processor and ocr_processor.available_openai_models:
                 openai_models = ocr_processor.available_openai_models
                 openai_dropdown_visible = True
                 openai_default_model = openai_models[0] # Default to the first available model
                 openai_status_text = f"✅ OpenAI Engine Available."
                 # Optionally list models in status: Models: {', '.join(openai_models)}
             else:
                 # Init succeeded but model fetch likely failed or returned empty
                 openai_status_text = "⚠️ OpenAI Engine Initialized, but failed to load models (Check Logs). Using default model."
                 openai_models = ["gpt-4o"] # Fallback list
                 openai_default_model = openai_models[0]
                 openai_dropdown_visible = True # Show dropdown even with default
        elif "OpenAI" in api_keys and api_keys["OpenAI"]:
            # Key provided, but engine not in available_engines (init failed)
            openai_status_text = "❌ OpenAI Engine Failed to Initialize. Check API Key/Logs."
        else:
             # No key provided or cleared
             openai_status_text = "ℹ️ Enter OpenAI API Key to enable."


        openai_model_dropdown_update = gr.Dropdown(
            choices=openai_models,
            value=openai_default_model,
            visible=openai_dropdown_visible,
            interactive=openai_dropdown_visible
        )

        # Return updates for all affected components in a dictionary
        return {
            available_engines_text: markdown_text,
            ocr_engine: radio_update,
            mistral_status: mistral_status_text,
            openai_status: openai_status_text,
            openai_model_select: openai_model_dropdown_update
        }

    # --- API key save/clear button logic --- #

    mistral_save_btn.click(
        fn=save_api_key,
        inputs=[mistral_key, gr.Text(value="Mistral", visible=False)],
        outputs=[mistral_status] # Let save_api_key return the immediate status
    ).then(
        fn=update_all_ui_elements, # Update all UI based on the new state
        outputs=[available_engines_text, ocr_engine, mistral_status, openai_status, openai_model_select]
    ).then(
        fn=lambda: "", # Clear input
        outputs=[mistral_key]
    )

    mistral_clear_btn.click(
        fn=clear_api_key,
        inputs=[gr.Text(value="Mistral", visible=False)],
        outputs=[mistral_status]
    ).then(
        fn=update_all_ui_elements,
        outputs=[available_engines_text, ocr_engine, mistral_status, openai_status, openai_model_select]
    ).then(
        fn=lambda: "",
        outputs=[mistral_key]
    )

    openai_save_btn.click(
        fn=save_api_key,
        inputs=[openai_key, gr.Text(value="OpenAI", visible=False)],
        outputs=[openai_status] # Let save_api_key return the immediate status
    ).then(
        fn=update_all_ui_elements,
        outputs=[available_engines_text, ocr_engine, mistral_status, openai_status, openai_model_select]
    ).then(
        fn=lambda: "",  # Clear input
        outputs=[openai_key]
    )

    openai_clear_btn.click(
        fn=clear_api_key,
        inputs=[gr.Text(value="OpenAI", visible=False)],
        outputs=[openai_status]
    ).then(
        fn=update_all_ui_elements,
        outputs=[available_engines_text, ocr_engine, mistral_status, openai_status, openai_model_select]
    ).then(
        fn=lambda: "",
        outputs=[openai_key]
    )

    # --- Clear Confirmation Handlers (Need update for new state/UI) --- #
    clear_btn.click(
        fn=handle_clear_click,
        inputs=[processed_results_state], # Check state instead of download_output
        outputs=[
            # Main UI elements to potentially clear or update
            file_input, ocr_engine, image_output, md_output,
            result_selector, download_format, download_selected_btn, download_all_btn,
            single_download_trigger, zip_download_trigger,
            processed_results_state, # Pass state through
            # Confirmation UI elements
            clear_confirm_msg, confirm_clear_btn, cancel_clear_btn,
            # Add the new markdown components
            download_options_md, download_trigger_md
        ]
    )
    confirm_clear_btn.click(
        fn=clear_and_hide_confirmation,
        inputs=[processed_results_state], # Pass state to the clearing function
        outputs=[
             # Main UI elements to clear
            file_input, ocr_engine, image_output, md_output,
            result_selector, download_format, download_selected_btn, download_all_btn,
            # Confirmation UI elements (to hide)
            clear_confirm_msg, confirm_clear_btn, cancel_clear_btn,
             # Trigger UI elements (to hide)
            single_download_trigger, zip_download_trigger,
            # Add the new markdown components
            download_options_md, download_trigger_md,
            processed_results_state # State is cleared and returned
        ]
    )
    cancel_clear_btn.click(
        fn=hide_confirmation,
        inputs=[],
        outputs=[
            clear_confirm_msg, confirm_clear_btn, cancel_clear_btn
        ]
    )

    # Add a startup event to initialize the UI correctly based on initial state
    demo.load(
        fn=update_all_ui_elements,
        outputs=[available_engines_text, ocr_engine, mistral_status, openai_status, openai_model_select]
    )

if __name__ == "__main__":
    auth_creds = (USERNAME, PASSWORD) if USERNAME and PASSWORD else None
    # Simplified launch logic
    demo.launch(auth=auth_creds, ssl_verify=True, share=False) # Set share=True if needed
    # demo.launch(ssl_verify=True) # Set share=True if needed 