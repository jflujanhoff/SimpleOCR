import gradio as gr
import os
import secrets
import zipfile
import shutil
from pathlib import Path
import logging
from docx import Document # Assuming python-docx is installed or added to requirements
from docx.shared import Inches # Might be useful for image insertion later

# Import MAX_FILES
from variables import MAX_FILES

logger = logging.getLogger(__name__)

class ProcessOCR:
    def __init__(self, ocr_processor=None, available_engines=None, ui_components=None):
        """Initialize the ProcessOCR interface handler."""
        # Store references if needed, or perform other setup
        self.ocr_processor = ocr_processor
        self.available_engines = available_engines
        self.ui_components = ui_components
        # TODO: Consider if output_dir should be passed here or fetched from ocr_processor
        self.output_dir = getattr(ocr_processor, 'output_dir', None)
        if not self.output_dir:
            logger.warning("ProcessOCR initialized without a valid output directory from ocr_processor.")
        logger.info("ProcessOCR class initialized.")


    # --- Helper function to clear output directory ---
    def _clear_output_directory(self, output_dir_path_str=None):
        """Removes all files and subdirectories within the specified directory."""
        # Use instance output_dir if available and no specific path is given
        if output_dir_path_str is None:
            output_dir_path_str = self.output_dir

        if not output_dir_path_str:
            logger.warning("Output directory path is not set. Skipping cleanup.")
            return

        output_dir_path = Path(output_dir_path_str)
        logger.info(f"Attempting to clear output directory: {output_dir_path}")

        if output_dir_path.exists() and output_dir_path.is_dir():
            for item in output_dir_path.iterdir():
                try:
                    if item.is_file() or item.is_symlink():
                        item.unlink()
                        logger.debug(f"Deleted file: {item}")
                    elif item.is_dir():
                        shutil.rmtree(item)
                        logger.debug(f"Deleted directory: {item}")
                except Exception as e:
                    logger.error(f"Error deleting item {item}: {e}", exc_info=True)
            logger.info(f"Successfully cleared output directory: {output_dir_path}")
        elif not output_dir_path.exists():
            logger.info(f"Output directory {output_dir_path} does not exist. No cleanup needed.")
        else:
            logger.warning(f"Path {output_dir_path} exists but is not a directory. Skipping cleanup.")

    # --- Helper function to create docx (text only for now) ---
    def _create_docx_file(self, text_content, output_path):
        """Creates a .docx file with the given text content."""
        try:
            document = Document()
            document.add_paragraph(text_content)
            # Future Enhancement: Could add logic here to insert images
            # image_paths = ... # Get image paths corresponding to text_content
            # for img_path in image_paths:
            #     if Path(img_path).exists():
            #         try:
            #             document.add_picture(img_path, width=Inches(6.0)) # Example width
            #         except Exception as img_e:
            #             logger.warning(f"Could not add image {img_path} to docx: {img_e}")
            document.save(output_path)
            logger.info(f"Successfully created DOCX file: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error creating DOCX file at {output_path}: {e}", exc_info=True)
            return False

    # --- NEW Helper function to write simple text files ---
    def _write_text_file(self, output_dir_path, base_filename, extension, text_content):
        """Writes text content to a file with the specified extension."""
        try:
            output_filename = f"{base_filename}.{extension}"
            output_path = output_dir_path / output_filename
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text_content)
            logger.info(f"Successfully wrote text file: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error writing {extension} file {output_filename} in {output_dir_path}: {e}", exc_info=True)
            return None

    # --- Helper function to prepare output file based on format ---
    def _prepare_output_file(self, base_filename, text_content, format_type, output_dir_path):
        """Creates a file (.txt, .md, .docx) in the specified directory."""
        output_path = None
        try:
            if format_type == "txt":
                # Use the new helper
                output_path = self._write_text_file(output_dir_path, base_filename, "txt", text_content)
                # No specific log here as it's in the helper

            elif format_type == "md":
                # Use the new helper
                output_path = self._write_text_file(output_dir_path, base_filename, "md", text_content)
                 # No specific log here as it's in the helper
                # OLD CODE:
                # output_filename = f"{base_filename}.md"
                # output_path = output_dir_path / output_filename
                # with open(output_path, "w", encoding="utf-8") as f:
                #     f.write(text_content) # For now, same as txt
                # logger.info(f"Prepared MD file: {output_path}")

            elif format_type == "doc": # Assuming 'doc' means '.docx'
                output_filename = f"{base_filename}.docx"
                output_path = output_dir_path / output_filename
                if not self._create_docx_file(text_content, output_path):
                    logger.error(f"Error creating DOCX file at {output_path}")
                    return None
                logger.info(f"Prepared DOCX file: {output_path}")

            else:
                logger.error(f"Unsupported format: {format_type}")
                return None

            return output_path

        except Exception as e:
            logger.error(f"Error preparing output file for {base_filename} as {format_type}: {e}", exc_info=True)
            return None

    # --- Main OCR Processing Function --- #
    # TODO: Refactor to use self.ocr_processor, self.available_engines, self.ui_components if initialized
    def process_document(self, files, ocr_engine, selected_openai_model, current_results_state, ocr_processor=None, available_engines=None, ui_components=None):
        """Process a list of documents, update state, and UI components."""
        # Use instance attributes if arguments are not provided (prefer arguments if passed)
        ocr_processor = ocr_processor or self.ocr_processor
        available_engines = available_engines or self.available_engines
        ui_components = ui_components or self.ui_components

        # Basic check if essential components are available
        if not ocr_processor or not available_engines or not ui_components:
             logger.error("ProcessOCR.process_document called without necessary components (ocr_processor, available_engines, ui_components).")
             # Return error state for all UI components
             num_outputs = 11 # Count based on the return signature
             error_md = "Error: Interface components not initialized. Check server logs."
             # Need to construct the return tuple matching the expected output structure
             # This assumes ui_components has the expected keys, otherwise this will fail
             error_updates = {
                 ui_components["result_selector"]: gr.Dropdown(choices=[], value=None, visible=False),
                 ui_components["md_output"]: error_md,
                 ui_components["image_output"]: gr.update(value=None, visible=False),
                 ui_components["download_format"]: gr.Radio(visible=False),
                 ui_components["download_selected_btn"]: gr.Button(visible=False),
                 ui_components["download_all_btn"]: gr.Button(visible=False),
                 ui_components["download_options_md"]: gr.update(visible=False),
                 ui_components["single_download_trigger"]: gr.update(value=None, visible=False),
                 ui_components["zip_download_trigger"]: gr.update(value=None, visible=False)
             }
             # Check if current_results_state is passed correctly
             if isinstance(current_results_state, dict):
                 return (*error_updates.values(), current_results_state)
             else:
                 # Fallback if state is not as expected
                 logger.error("current_results_state is not a dict in process_document error path.")
                 return (*error_updates.values(), {"text": {}, "images": {}})


        logger.info(f"Received {len(files) if files else 0} file(s) for processing with engine {ocr_engine}")

        # --- Initialize UI update dictionary --- #
        initial_updates = {
            ui_components["result_selector"]: gr.Dropdown(choices=[], value=None, visible=False),
            ui_components["md_output"]: "Initializing...", # Changed initial message
            ui_components["image_output"]: gr.update(value=None, visible=False),
            ui_components["download_format"]: gr.Radio(visible=False),
            ui_components["download_selected_btn"]: gr.Button(visible=False),
            ui_components["download_all_btn"]: gr.Button(visible=False),
            ui_components["download_options_md"]: gr.update(visible=False),
            ui_components["single_download_trigger"]: gr.update(value=None, visible=False),
            ui_components["zip_download_trigger"]: gr.update(value=None, visible=False)
        }

        # --- Input Validation --- #
        if not files:
            initial_updates[ui_components["md_output"]] = "Error: No files uploaded."
            # Return the current state if no files
            return (*initial_updates.values(), current_results_state)

        # --- Check File Count Limit --- #
        if len(files) > MAX_FILES:
            error_msg = f"Error: Too many files uploaded. Maximum allowed is {MAX_FILES}. You uploaded {len(files)}."
            logger.warning(error_msg)
            gr.Warning(error_msg) # Display warning popup to user
            initial_updates[ui_components["md_output"]] = error_msg # Update markdown as well
            # Return the current state if too many files
            return (*initial_updates.values(), current_results_state)

        # --- Clear Output Directory Before Processing --- #
        # Use the instance's output_dir if available
        output_dir_to_clear = getattr(ocr_processor, 'output_dir', self.output_dir)
        if output_dir_to_clear:
             self._clear_output_directory(output_dir_to_clear)
        else:
            logger.warning("Output directory not available, cannot clear output directory.")


        # --- Initialize counts for successful files --- #
        processed_results = {"text": {}, "images": {}}
        errors_occurred = False
        error_messages = []
        successful_file_count = 0 # Count only non-error files
        total_page_count = 0      # Sum pages only from non-error files

        # --- Pre-processing Checks --- #
        # Note: ocr_processor should be initialized before calling this method
        if ocr_processor is None:
            initial_updates[ui_components["md_output"]] = "Error: OCR processor is not initialized. Check server logs."
            logger.error("process_document called with uninitialized OCR processor.")
            return (*initial_updates.values(), current_results_state)

        if ocr_engine not in available_engines:
            initial_updates[ui_components["md_output"]] = f"Error: {ocr_engine} OCR is not available. Check API keys or server logs."
            return (*initial_updates.values(), current_results_state)

        if ocr_engine == "OpenAI" and not selected_openai_model:
            if not (ocr_processor and ocr_processor.available_openai_models):
                 initial_updates[ui_components["md_output"]] = "Error: OpenAI engine selected, but no models could be loaded."
            else:
                initial_updates[ui_components["md_output"]] = "Error: OpenAI engine selected, but no specific model chosen."
            return (*initial_updates.values(), current_results_state)

        # --- Process Each File --- #
        for file_obj in files:
            # Gradio File objects might have absolute paths, ensure we only use the filename part
            original_filename = Path(getattr(file_obj, 'name', f"unknown_file_{secrets.token_hex(4)}")).name
            logger.info(f"Processing file: {original_filename}")

            # --- Read file content immediately --- #
            file_content = None
            error_reading_file = None
            try:
                file_path = file_obj.name
                logger.info(f"Attempting to read temporary file path: {file_path} for original: {original_filename}")
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                if not file_content:
                    error_reading_file = f"Error: File '{original_filename}' is empty."
                    logger.error(error_reading_file)
                else:
                    logger.info(f"Successfully read {len(file_content)} bytes from temp file for {original_filename}")
            except Exception as read_e:
                error_reading_file = f"Error reading temporary file for {original_filename}: {read_e}"
                logger.error(error_reading_file, exc_info=True)
            # --- End read --- #

            # --- Skip if read failed --- #
            if error_reading_file:
                errors_occurred = True
                error_messages.append(f"- {original_filename}: Failed to read file content.")
                processed_results["text"][original_filename] = error_reading_file # Store read error
                processed_results["images"][original_filename] = []
                continue # Skip to the next file
            # --- End skip --- #

            try:
                # ... (logging before call - can be removed if desired) ...
                # logger.info(f"[process_document callback] Before calling process_document for {original_filename}:")
                # logger.info(f"  ocr_processor object: {ocr_processor}")
                # mistral_engine_state = getattr(ocr_processor, 'mistral', 'Attribute not found')
                # logger.info(f"  ocr_processor.mistral state: {mistral_engine_state}")

                # --- Pass file_content and filename, not file_obj --- #
                _ , image_paths, result_text = ocr_processor.process_document(
                    file_content=file_content,
                    file_name=original_filename, # Pass the extracted filename
                    ocr_engine=ocr_engine,
                    openai_model=selected_openai_model if ocr_engine == "OpenAI" else None
                )
                # --- End change --- #

                # --- Handle Results --- #
                if result_text is not None and result_text.startswith("Error:"):
                    # Handle Error Result
                    logger.error(f"Error processing {original_filename}: {result_text}")
                    errors_occurred = True
                    error_messages.append(f"- {original_filename}: {result_text}")
                    processed_results["text"][original_filename] = result_text
                    processed_results["images"][original_filename] = image_paths or []
                elif result_text is None:
                    # Handle Null Result (treat as error)
                    logger.warning(f"No text extracted from {original_filename}.")
                    errors_occurred = True
                    error_msg = f"- {original_filename}: Could not extract text."
                    error_messages.append(error_msg)
                    processed_results["text"][original_filename] = "Error: Could not extract text."
                    processed_results["images"][original_filename] = image_paths or []
                else:
                    # Handle Success Result
                    logger.info(f"Successfully processed {original_filename}.")
                    processed_results["text"][original_filename] = result_text
                    processed_results["images"][original_filename] = image_paths or []

                    # --- Increment counts only on success --- #
                    successful_file_count += 1
                    # --- Count pages based on number of returned image paths --- #
                    pages_in_file = len(image_paths) if image_paths else 0
                    total_page_count += pages_in_file
                    logger.debug(f"Counted {pages_in_file} pages for {original_filename} based on image paths.")

            except Exception as e:
                # Handle Exception during processing
                error_msg = f"Error processing file {original_filename}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                errors_occurred = True
                error_messages.append(f"- {original_filename}: Processing failed unexpectedly. Check logs.")
                processed_results["text"][original_filename] = f"Error: Processing failed unexpectedly."
                processed_results["images"][original_filename] = []

        # --- Final UI Update Preparation --- #

        # Calculate count string using counters updated during the loop
        count_string = f"Processed Files: {successful_file_count} | Total Pages: {total_page_count}"
        logger.info(f"Final counts (successful files only): {count_string}")

        # --- Determine initial display text (errors or first result) --- #
        if errors_occurred:
            # Prepend error messages to the first successful result's text if any exist
            # Or display errors in a dedicated area (better UX, but requires UI change)
            final_md_output = "\\n---\\n".join(error_messages)
            # If there are also results, show them after errors
            if processed_results["text"]:
                first_filename = list(processed_results["text"].keys())[0]
                final_md_output += "\\n\\n**First Processed Result:**\\n" + "\\n".join(processed_results["text"][first_filename]) # Assuming list format
            logger.warning(f"Processing finished with errors: {error_messages}")
            # Display the first error in the MD output for immediate feedback
            # updates[ui_components["md_output"]] = error_messages[0] # Replaced by final_md_output below
        elif not processed_results["text"]:
             # No errors, but also no results (e.g., empty files?)
             final_md_output = "Processing complete, but no text could be extracted."
             logger.info("Processing complete, but no text was extracted.")
        else:
            # Display the first result
            first_filename = list(processed_results["text"].keys())[0]
            # Assign the text directly, don't join characters
            final_md_output = processed_results["text"][first_filename]
            logger.info(f"Processing successful. Displaying first result: {first_filename}")

        # --- Prepare Final Updates Dictionary --- #
        # Start with initial (mostly hidden) state
        updates = initial_updates.copy()
        # Override specific components based on results

        # Populate dropdown and potentially show first result
        if processed_results["text"]:
            filenames = list(processed_results["text"].keys())
            first_filename = filenames[0]
            updates[ui_components["result_selector"]] = gr.Dropdown(choices=filenames, value=first_filename, label="Select Processed File", visible=True, interactive=True)
            updates[ui_components["md_output"]] = final_md_output # Show first result or errors

            # Show images for the first result if available
            first_image_paths = processed_results.get("images", {}).get(first_filename, [])
            valid_first_image_paths = [p for p in first_image_paths if p and Path(p).exists()]
            if valid_first_image_paths:
                 updates[ui_components["image_output"]] = gr.update(value=valid_first_image_paths, visible=True)
            else:
                 updates[ui_components["image_output"]] = gr.update(value=None, visible=False) # Ensure it's hidden if no images


            # Make download section visible
            updates[ui_components["download_group"]] = gr.update(visible=True) # Show download group
            updates[ui_components["download_options_md"]] = gr.update(visible=True)
            updates[ui_components["download_format"]] = gr.Radio(choices=["txt", "md", "doc"], value="txt", label="Format", visible=True, interactive=True)
            updates[ui_components["download_selected_btn"]] = gr.Button(visible=True, interactive=True)
            updates[ui_components["download_all_btn"]] = gr.Button(visible=True, interactive=True)
            # Keep download triggers hidden, they are activated by button clicks
            updates[ui_components["single_download_trigger"]] = gr.update(value=None, visible=False, interactive=False)
            updates[ui_components["zip_download_trigger"]] = gr.update(value=None, visible=False, interactive=False)
            updates[ui_components["result_group"]] = gr.update(visible=True) # Make result group visible


        else:
            # No results, keep things hidden/disabled
            updates[ui_components["result_selector"]] = gr.Dropdown(choices=[], value=None, label="Select Processed File", visible=False, interactive=False)
            updates[ui_components["md_output"]] = final_md_output # Show "No text extracted" or errors
            updates[ui_components["image_output"]] = gr.update(value=None, visible=False)
            updates[ui_components["download_group"]] = gr.update(visible=False) # Hide download group
            updates[ui_components["download_options_md"]] = gr.update(visible=False)
            updates[ui_components["download_format"]] = gr.Radio(visible=False, interactive=False)
            updates[ui_components["download_selected_btn"]] = gr.Button(visible=False)
            updates[ui_components["download_all_btn"]] = gr.Button(visible=False)
            updates[ui_components["result_group"]] = gr.update(visible=True) # Still show result group, but it will be mostly empty/disabled


        # The order MUST match the 'outputs_process' list in interface.py
        return (
            updates[ui_components["result_group"]],
            updates[ui_components["result_selector"]],
            updates[ui_components["md_output"]],
            updates[ui_components["image_output"]],
            updates[ui_components["download_group"]],
            updates[ui_components["download_format"]],
            updates[ui_components["download_selected_btn"]],
            updates[ui_components["download_all_btn"]],
            updates[ui_components["download_options_md"]],
            updates[ui_components["single_download_trigger"]],
            updates[ui_components["zip_download_trigger"]],
            processed_results,  # Return the full state
            count_string # Use the correctly calculated count string
        )


    # --- Download Handlers --- #
    def download_selected_file(self, selected_filename, format_type, current_results_state, ocr_processor=None):
        """Prepares a single file for download based on selected format."""
        ocr_processor = ocr_processor or self.ocr_processor
        output_dir = getattr(ocr_processor, 'output_dir', self.output_dir)
        # Default return value for failure/skip cases
        fail_return = (gr.update(value=None, visible=False), gr.update(visible=False))

        if not output_dir:
            logger.error("Output directory not configured for download.")
            gr.Error("Server configuration error: Output directory not set.")
            return fail_return # Return tuple

        if not selected_filename or not current_results_state or selected_filename not in current_results_state.get("text", {}):
            logger.warning(f"Download failed: Selected filename '{selected_filename}' not found in results state.")
            gr.Warning(f"Could not find result for '{selected_filename}'. Please process files first.")
            return fail_return # Return tuple

        text_content = current_results_state["text"][selected_filename]
        # Skip download if the result was an error message
        if text_content.startswith("Error:"):
             logger.warning(f"Download skipped: Result for '{selected_filename}' is an error.")
             gr.Info(f"Cannot download '{selected_filename}' as it contains an error message.")
             return fail_return # Return tuple

        # Base filename without extension
        base_filename = Path(selected_filename).stem
        # Define output path within the configured output directory
        output_dir_path = Path(output_dir)
        output_dir_path.mkdir(parents=True, exist_ok=True) # Ensure directory exists

        try:
            output_path = self._prepare_output_file(base_filename, text_content, format_type, output_dir_path)

            # If we reach here and output_path is set, it means success
            if output_path and output_path.exists():
                return (
                    gr.update(value=str(output_path), visible=True),
                    gr.update(visible=True) # Make trigger group visible
                )
            else:
                 # Should not happen if logic is correct, but as a fallback
                 logger.error(f"File creation failed unexpectedly for {selected_filename} with format {format_type}")
                 gr.Error(f"Failed to prepare download for {selected_filename}.")
                 return fail_return

        except Exception as e:
            logger.error(f"Error preparing file '{selected_filename}' for download as {format_type}: {e}", exc_info=True)
            gr.Error(f"Failed to prepare download for {selected_filename}.")
            return fail_return # Return tuple

    def download_all_files(self, format_type, current_results_state, ocr_processor=None):
        """Prepares a ZIP archive containing all processed files in the selected format."""
        ocr_processor = ocr_processor or self.ocr_processor
        output_dir = getattr(ocr_processor, 'output_dir', self.output_dir)
        # Default return value for failure/skip cases
        fail_return = (gr.update(value=None, visible=False), gr.update(visible=False))

        if not output_dir:
            logger.error("Output directory not configured for ZIP download.")
            gr.Error("Server configuration error: Output directory not set.")
            return fail_return # Return tuple

        if not current_results_state or not current_results_state.get("text"):
            logger.warning("Download All failed: No results found in state.")
            gr.Info("No processed files available to download.")
            return fail_return # Return tuple

        output_dir_path = Path(output_dir)
        # Create a unique sub-directory for temporary files to avoid conflicts
        temp_zip_dir = output_dir_path / f"temp_zip_{secrets.token_hex(8)}"
        zip_filepath = None # Define zip_filepath outside try block
        try:
            temp_zip_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"Created temporary directory for zipping: {temp_zip_dir}")

            zip_filename = f"ocr_results_{format_type}_{secrets.token_hex(4)}.zip"
            zip_filepath = output_dir_path / zip_filename
            files_added_count = 0

            with zipfile.ZipFile(zip_filepath, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for original_filename, text_content in current_results_state["text"].items():
                    # Skip files that resulted in errors
                    if text_content.startswith("Error:"):
                        logger.debug(f"Skipping error result for {original_filename} in zip.")
                        continue

                    base_filename = Path(original_filename).stem
                    # temp_file_path = None # Path for the intermediate file - Handled by helper
                    # temp_filename = None # Define temp_filename here - No longer needed here

                    try:
                        temp_file_path = self._prepare_output_file(base_filename, text_content, format_type, temp_zip_dir)

                        # Add the created file to the zip
                        # Check if the helper returned a valid path and the file exists
                        if temp_file_path and temp_file_path.exists():
                            # Use the file's actual name for the archive name
                            zipf.write(temp_file_path, arcname=temp_file_path.name)
                            logger.debug(f"Added {temp_file_path.name} to zip.")
                            files_added_count += 1
                        else:
                            # Log if the helper failed or the file doesn't exist
                            logger.warning(f"File preparation failed for {original_filename} (format: {format_type}). Not added to zip.")

                    except Exception as file_e:
                        logger.error(f"Error processing file {original_filename} for zip archive: {file_e}", exc_info=True)

            if files_added_count > 0:
                logger.info(f"Created ZIP archive with {files_added_count} file(s): {zip_filepath}")
                return (
                    gr.update(value=str(zip_filepath), visible=True),
                    gr.update(visible=True) # Make trigger group visible
                )
            else:
                logger.warning("No valid files were added to the ZIP archive.")
                gr.Info("No valid results available to include in the download.")
                # Clean up empty zip file if created
                if zip_filepath and zip_filepath.exists():
                    try:
                        zip_filepath.unlink()
                    except OSError as del_e:
                        logger.error(f"Error deleting empty zip file {zip_filepath}: {del_e}")
                return fail_return # Return tuple

        except Exception as e:
            logger.error(f"Error creating ZIP archive: {e}", exc_info=True)
            gr.Error("Failed to create ZIP archive.")
            return fail_return # Return tuple
        finally:
            # Clean up the temporary directory
            if temp_zip_dir.exists():
                try:
                    shutil.rmtree(temp_zip_dir)
                    logger.info(f"Cleaned up temporary zip directory: {temp_zip_dir}")
                except Exception as cleanup_e:
                    logger.error(f"Error cleaning up temporary zip directory {temp_zip_dir}: {cleanup_e}", exc_info=True)


# --- Standalone UI Update Functions (Consider moving to UIInteractions or keeping separate) ---
# These were likely intended to be called directly if ProcessOCR wasn't handling UI updates,
# but currently, process_document returns all necessary UI updates.
# Keeping them commented out or removing them might be cleaner unless used elsewhere.

# def update_output_display(selected_filename, current_results_state):
#     """Updates the markdown and image preview based on dropdown selection."""
#     # This logic is now handled within UIInteractions.display_selected_result
#     pass

# def get_ui_components_dict(*args):
#     """Helper to create a dictionary from UI component references."""
#     # This seems redundant if ui_components is managed correctly in interface.py
#     pass

# Example of how clear_output_directory might be used if passed separately
# def clear_output_directory(output_dir):
#     path = Path(output_dir)
#     if path.exists() and path.is_dir():
#         logger.info(f"Clearing output directory: {output_dir}")
#         # ... (implementation as in _clear_output_directory)
#     else:
#         logger.warning(f"Cannot clear non-existent or non-directory path: {output_dir}")
